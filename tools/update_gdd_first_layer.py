from pathlib import Path

from docx import Document

from build_gdd_docx import patch_docx_fonts


DOCX = Path(__file__).resolve().parents[1] / "心智后台_完整游戏策划案.docx"


def replace_paragraphs(doc: Document) -> None:
    replacements = {
        "第一层应让玩家快速理解系统，并通过简单 Bug 获得正反馈。":
            "第一层应让玩家快速理解系统，并通过 Bug、错字虫、红点提醒等怪物获得基础战斗反馈。",
        "SyntaxError": "红字审判",
        "????": "红字审判",
        "定位：语法错误的具象化，用于检验玩家是否理解基本攻防和手牌管理。":
            "定位：第一层综合测试 Boss，用更高血量、护盾和手牌干扰检验玩家是否理解基本攻防、参数分配和回合规划。",
        "周期性删除玩家一行代码。": "周期性干扰玩家手牌。",
        "设计目的：让玩家意识到程序区不是无风险堆叠，需要提前完成关键输出。":
            "设计目的：让玩家意识到第一层已经需要观察敌方代码，不能只堆攻击，也要安排防御、保留和下回合资源。",
        "普通战斗平均 2 至 4 回合结束。":
            "普通战斗平均 3 至 4 回合结束，避免怪物过快死亡导致机制来不及出现。",
        "Boss 允许玩家犯 1 至 2 次轻微错误。":
            "第一层 Boss 目标战斗长度为 6 至 8 回合，允许玩家犯 1 至 2 次轻微错误。",
        "调整第一层敌人伤害，保证新手不被早期惩罚过重。":
            "继续观察第一层敌人的 HP、攻击和干扰频率，保证机制足够明显但不形成早期硬卡点。",
        "当前版本缺少完整教学。建议以第一场 Bug 战作为强制教学。":
            "当前版本缺少完整教学。建议以第一场 Bug 战作为强制教学。",
        "进入战斗后高亮敌人意图，提示“Bug 本回合将攻击 5 点”。":
            "进入战斗后高亮敌人意图，提示“Bug 本回合将攻击 7 点”。",
    }

    for para in doc.paragraphs:
        if para.text in replacements:
            para.text = replacements[para.text]

    # Repair known locations in case an earlier script wrote question marks.
    if len(doc.paragraphs) > 62:
        doc.paragraphs[62].text = "第一层应让玩家快速理解系统，并通过 Bug、错字虫、红点提醒等怪物获得基础战斗反馈。"
    if len(doc.paragraphs) > 188:
        doc.paragraphs[188].text = "红字审判"
    if len(doc.paragraphs) > 190:
        doc.paragraphs[190].text = "定位：第一层综合测试 Boss，用更高血量、护盾和手牌干扰检验玩家是否理解基本攻防、参数分配和回合规划。"
    if len(doc.paragraphs) > 194:
        doc.paragraphs[194].text = "周期性干扰玩家手牌。"
    if len(doc.paragraphs) > 195:
        doc.paragraphs[195].text = "设计目的：让玩家意识到第一层已经需要观察敌方代码，不能只堆攻击，也要安排防御、保留和下回合资源。"
    if len(doc.paragraphs) > 277:
        doc.paragraphs[277].text = "普通战斗平均 3 至 4 回合结束，避免怪物过快死亡导致机制来不及出现。"
    if len(doc.paragraphs) > 278:
        doc.paragraphs[278].text = "第一层 Boss 目标战斗长度为 6 至 8 回合，允许玩家犯 1 至 2 次轻微错误。"
    if len(doc.paragraphs) > 296:
        doc.paragraphs[296].text = "继续观察第一层敌人的 HP、攻击和干扰频率，保证机制足够明显但不形成早期硬卡点。"
    if len(doc.paragraphs) > 378:
        doc.paragraphs[378].text = "当前版本缺少完整教学。建议以第一场 Bug 战作为强制教学。"
    if len(doc.paragraphs) > 388:
        doc.paragraphs[388].text = "进入战斗后高亮敌人意图，提示“Bug 本回合将攻击 7 点”。"


def update_tables(doc: Document) -> None:
    t1 = doc.tables[1]
    t1.rows[1].cells[1].text = "Bug 领域"
    t1.rows[1].cells[2].text = "日常焦虑、弹窗打断、低级错误与临时催促"
    t1.rows[1].cells[3].text = "红字审判"

    enemy_rows = [
        ["敌人", "层级", "行为特点"],
        ["Bug", "第一层", "代码库为攻击 7、防御 7；行动循环为攻击 -> 防御 -> 攻击+防御；当前 HP 32"],
        ["错字虫", "第一层", "小范围随机攻击；当前 HP 30，攻击 6-9"],
        ["红点提醒", "第一层", "每 3 回合触发提醒，下回合少抽 1 张牌"],
        ["空值幽影", "第一层/精英", "潜伏一回合后打出较高伤害，形成节奏压力"],
        ["催单员", "第一层/精英", "伤害随回合递增，拖久会变危险"],
        ["待办残片", "第一层/精英", "攻击和防御交替出现，教学玩家观察敌方代码"],
        ["InfiniteLoop", "第二层", "伤害随回合递增"],
        ["MemoryLeak", "第二层", "吸血攻击"],
        ["RaceCondition", "第二/三层", "高波动随机攻击"],
        ["GarbageCollector", "第二层", "干扰手牌"],
        ["Recursion", "第二/三层", "指数级伤害成长"],
        ["StackOverflow", "第二/三层", "多段循环攻击"],
        ["Deadlock", "第二/三层", "防御与爆发交替"],
    ]

    t7 = doc.tables[7]
    while len(t7.rows) < len(enemy_rows):
        t7.add_row()
    for row_idx, values in enumerate(enemy_rows):
        for col_idx, value in enumerate(values):
            t7.rows[row_idx].cells[col_idx].text = value
    for row_idx in range(len(enemy_rows), len(t7.rows)):
        for cell in t7.rows[row_idx].cells:
            cell.text = ""


def main() -> None:
    doc = Document(DOCX)
    replace_paragraphs(doc)
    update_tables(doc)
    doc.save(DOCX)
    patch_docx_fonts(DOCX)
    print(f"updated {DOCX.name}")


if __name__ == "__main__":
    main()
