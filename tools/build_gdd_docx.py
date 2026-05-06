from __future__ import annotations

import re
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "心智后台_完整游戏策划案.md"
OUT = ROOT / "心智后台_完整游戏策划案.docx"

BLACK = RGBColor(0, 0, 0)
LIGHT_FILL = "FFFFFF"
BORDER = "BFBFBF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_cm) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK


def set_run_font(run, size=None, bold=None, color=BLACK, name="宋体") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = color


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:cs"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    title_style = styles["Title"]
    title_style.font.name = "黑体"
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "黑体")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "黑体")
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title_style._element.rPr.rFonts.set(qn("w:cs"), "黑体")
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = BLACK

    for name, size, before, after in (
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 8, 4),
        ("Heading 3", 12, 8, 4),
    ):
        st = styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:ascii"), "黑体")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "黑体")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st._element.rPr.rFonts.set(qn("w:cs"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BLACK
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "宋体"
        st._element.rPr.rFonts.set(qn("w:ascii"), "宋体")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "宋体")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        st._element.rPr.rFonts.set(qn("w:cs"), "宋体")
        st.font.size = Pt(12)
        st.font.color.rgb = BLACK
        st.paragraph_format.space_after = Pt(4)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc: Document, title: str, meta_lines: list[str]) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(90)
    r = p.add_run(title.replace("# ", ""))
    set_run_font(r, size=22, bold=True, color=BLACK, name="黑体")

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(18)
    r = sub.add_run("普通策划案")
    set_run_font(r, size=12, color=BLACK, name="宋体")

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_borders(table)
    set_column_widths(table, [3.6, 12.2])
    for raw in meta_lines:
        if "：" in raw:
            key, val = raw.split("：", 1)
        else:
            key, val = "说明", raw
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = val.strip()
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=12, bold=(cell is row.cells[0]), color=BLACK, name="宋体")

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(20)
    note.paragraph_format.space_after = Pt(0)
    r = note.add_run("使用说明")
    set_run_font(r, size=12, bold=True, color=BLACK, name="黑体")
    p = doc.add_paragraph(
        "本文档按当前可玩原型整理，用于记录玩法思路、系统拆解、内容补全与持续迭代。"
        "后续修改建议继续保留章节编号与表格结构，便于追踪需求变化。"
    )
    p.paragraph_format.left_indent = Cm(0.2)
    doc.add_page_break()


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        parts = [p.strip() for p in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts):
            rows.append(parts)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)

    usable_cm = 17.4
    if col_count == 2:
        widths = [4.3, usable_cm - 4.3]
    elif col_count == 3:
        widths = [3.3, 4.5, usable_cm - 7.8]
    elif col_count == 4:
        widths = [2.4, 3.1, 6.2, usable_cm - 11.7]
    else:
        widths = [usable_cm / col_count] * col_count

    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        if ridx == 0:
            set_repeat_table_header(row)
        for cidx in range(col_count):
            cell = row.cells[cidx]
            cell.text = row_data[cidx] if cidx < len(row_data) else ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if ridx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=12, bold=(ridx == 0), color=BLACK, name="宋体")
    set_column_widths(table, widths)
    doc.add_paragraph()


def add_workflow_table(doc: Document) -> None:
    rows = [
        ["阶段", "玩家动作", "系统反馈"],
        ["1", "标题界面选择“开始调试”", "进入角色选择"],
        ["2", "选择人格模块", "生成初始牌组、地图与局内状态"],
        ["3", "在地图选择节点", "进入战斗、事件、商店、休息或宝箱"],
        ["4", "战斗中编排程序并运行", "根据执行结果进入奖励、失败或继续回合"],
        ["5", "击败本层 Boss", "进入下一层；最终层 Boss 击败后胜利结算"],
    ]
    add_table(doc, rows)


def force_paragraph_runs_font(paragraph, font_name: str, size: int, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=BLACK, name=font_name)


def add_paragraph_with_inline_code(doc: Document, text: str, style=None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            set_run_font(r, size=12, name="宋体", color=BLACK)
        else:
            r = p.add_run(part)
            set_run_font(r, size=12, name="宋体", color=BLACK)
    return p


def patch_docx_fonts(path: Path) -> None:
    """Remove theme font fallbacks that WPS can display as MS Gothic."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    q = lambda name: f"{{{w_ns}}}{name}"

    def set_rfonts(el, font: str) -> None:
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            el.attrib.pop(q(attr), None)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            el.set(q(attr), font)

    tmp = path.with_suffix(".tmp.docx")
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in {"word/styles.xml", "word/stylesWithEffects.xml", "word/document.xml"}:
                root = etree.fromstring(data)
                if item.filename in {"word/styles.xml", "word/stylesWithEffects.xml"}:
                    # Clear every theme font reference first.
                    for rf in root.xpath(".//w:rFonts", namespaces=ns):
                        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
                            rf.attrib.pop(q(attr), None)
                    heading_ids = {"Title", "TitleChar", "Heading1", "Heading1Char", "Heading2", "Heading2Char", "Heading3", "Heading3Char"}
                    body_ids = {"Normal", "BodyText", "ListBullet", "ListNumber"}
                    for style in root.xpath(".//w:style", namespaces=ns):
                        sid = style.get(q("styleId"))
                        rpr = style.find("w:rPr", namespaces=ns)
                        if rpr is None:
                            rpr = etree.SubElement(style, q("rPr"))
                        rf = rpr.find("w:rFonts", namespaces=ns)
                        if rf is None:
                            rf = etree.SubElement(rpr, q("rFonts"))
                        if sid in heading_ids:
                            set_rfonts(rf, "黑体")
                        elif sid in body_ids:
                            set_rfonts(rf, "宋体")
                else:
                    # Direct-format heading/title runs too, so the toolbar does not infer theme fonts.
                    for p in root.xpath(".//w:p", namespaces=ns):
                        pstyle = p.xpath("./w:pPr/w:pStyle/@w:val", namespaces=ns)
                        if not pstyle:
                            continue
                        font = "黑体" if pstyle[0] in {"Title", "Heading1", "Heading2", "Heading3"} else None
                        if not font:
                            continue
                        for r in p.xpath("./w:r", namespaces=ns):
                            rpr = r.find("w:rPr", namespaces=ns)
                            if rpr is None:
                                rpr = etree.Element(q("rPr"))
                                r.insert(0, rpr)
                            rf = rpr.find("w:rFonts", namespaces=ns)
                            if rf is None:
                                rf = etree.SubElement(rpr, q("rFonts"))
                            set_rfonts(rf, font)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            zout.writestr(item, data)
    tmp.replace(path)


def build_docx() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    title = lines[0].strip()
    meta = []
    idx = 1
    while idx < len(lines) and lines[idx].strip() != "---":
        if lines[idx].strip():
            meta.append(lines[idx].strip().rstrip("  "))
        idx += 1
    add_cover(doc, title, meta)

    idx += 1
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                if code_lang == "mermaid":
                    add_workflow_table(doc)
                else:
                    for cl in code_lines:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(0.5)
                        r = p.add_run(cl)
                        set_run_font(r, size=12, name="宋体", color=BLACK)
                in_code = False
                code_lang = ""
                code_lines = []
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            idx += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue

        if stripped.startswith("#### "):
            para = doc.add_heading(stripped[5:].strip(), level=2)
            force_paragraph_runs_font(para, "黑体", 12, True)
        elif stripped.startswith("### "):
            para = doc.add_heading(stripped[4:].strip(), level=2)
            force_paragraph_runs_font(para, "黑体", 12, True)
        elif stripped.startswith("## "):
            para = doc.add_heading(stripped[3:].strip(), level=1)
            force_paragraph_runs_font(para, "黑体", 14, True)
        elif stripped.startswith("# "):
            para = doc.add_heading(stripped[2:].strip(), level=1)
            force_paragraph_runs_font(para, "黑体", 14, True)
        elif stripped.startswith("- "):
            add_paragraph_with_inline_code(doc, stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            add_paragraph_with_inline_code(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        else:
            paragraph = add_paragraph_with_inline_code(doc, stripped)
            if stripped.endswith("：") and len(stripped) <= 18:
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.color.rgb = BLACK
                paragraph.runs[0].font.name = "黑体"
                paragraph.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        idx += 1

    doc.core_properties.title = "《心智后台》完整游戏策划案"
    doc.core_properties.subject = "普通策划案"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    patch_docx_fonts(OUT)


if __name__ == "__main__":
    build_docx()
